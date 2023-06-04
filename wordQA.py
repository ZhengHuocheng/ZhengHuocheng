from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import streamlit as st
# 定义格式检查规则
RULES = {
    '页面大小': ('Letter', 21, 29.7),  # 页面大小（纸张类型、宽度、高度cm）
    '页边距': {
        '左边距': 3.17,  # 单位：cm
        '右边距': 3.17,
        '上边距': 2.54,
        '下边距': 2.54
    },
    '题目': {
        '字体': '黑体',
        '字号': 16,
        '对齐方式': '居中'
    },
    'Heading 1': {
        '字体': '黑体',
        '字号': 10,
        '对齐方式': '左对齐',
    },
    'Heading 2': {
        '字体': '黑体',
        '字号': 12,
        '对齐方式': '左对齐',
    },
    'Heading 3': {
        '字体': '黑体',
        '字号': 12,
        '对齐方式': '左对齐',
    },
    '正文': {
        '字体': '宋体',
        '字号': 10,    #磅
        '首行缩进': 20,  #磅，缩进磅数/字号≈一个字符大小
        '行距': 12    #磅
    }
}

# 打开Word文档
doc = Document('other_train.docx')

# 错误信息列表
errors = []

# 提取页面大小
sections = doc.sections
if sections:
    first_section = sections[0]
    page_width = round(first_section.page_width.cm, 2)  # 将页面宽度转换为cm，并保留两位小数
    page_height = round(first_section.page_height.cm, 2)  # 将页面高度转换为cm，并保留两位小数
    expected_page_size = RULES['页面大小'][1:]
    if (page_width, page_height) != expected_page_size:
        error = f"页面大小为{page_width} x {page_height} cm，不符合规定的{expected_page_size[0]} x {expected_page_size[1]} cm"
        errors.append(error)

# 提取页边距
margins = (
    round(first_section.left_margin.cm, 2),  # 将左边距转换为cm，并保留两位小数
    round(first_section.right_margin.cm, 2),  # 将右边距转换为cm，并保留两位小数
    round(first_section.top_margin.cm, 2),  # 将上边距转换为cm，并保留两位小数
    round(first_section.bottom_margin.cm, 2)  # 将下边距转换为cm，并保留两位小数
)
expected_margins = RULES['页边距']
for i, (margin, expected_margin) in enumerate(zip(margins, expected_margins.values()), 1):
    if margin != expected_margin:
        error = f"第{i}个页边距为{margin} cm，不符合规定的{expected_margin} cm"
        errors.append(error)
title_index=[]
# 提取题目信息
title_paragraph = doc.paragraphs[0]
expected_title_info = RULES['题目']
title_run = title_paragraph.runs[0]
title_font_name = title_run.font.name
title_font_size = title_run.font.size.pt
title_index.append(0)
# 提取对齐方式
title_alignment = str(title_paragraph.alignment)
if title_alignment == 'CENTER (1)':
    title_alignment = '居中'
elif title_alignment == 'RIGHT (2)':
    title_alignment = '右对齐'
elif title_alignment == 'JUSTIFY (3)':
    title_alignment = '两端对齐'
elif title_alignment == 'LEFT (0)':
    title_alignment = '左对齐'
if (title_font_name, title_font_size, title_alignment) != (expected_title_info['字体'], expected_title_info['字号'], expected_title_info['对齐方式']):
    error = f"题目的字体为'{title_font_name}'，字号为{title_font_size}，对齐方式为'{title_alignment}'，不符合规定的'{expected_title_info['字形']}'、{expected_title_info['字号']}、'{expected_title_info['对齐方式']}'"
    errors.append(error)

# 提取各级标题信息
# 定义各级标题样式名称
heading_styles = ['Heading 1', 'Heading 2', 'Heading 3']
# 遍历文档中的段落
for i,paragraph in enumerate(doc.paragraphs):
    style_name = paragraph.style.name 
    # 判断段落的样式是否为各级标题样式
    if style_name in heading_styles:
        title_index.append(i)
        # 获取标题级别
        heading_level = heading_styles.index(style_name) + 1      
        # 输出标题级别和文本内容
        #print(f"Level {heading_level} Heading: {paragraph.text}")
        # 提取本次标题的字体、字号、对齐方式和行距
        this_title_paragraph = doc.paragraphs[i]
        this_title_run = this_title_paragraph.runs[0]
        this_title_font_name = this_title_run.font.name
        this_title_font_size = this_title_run.font.size.pt
        # 提取对齐方式
        this_title_alignment = str(this_title_paragraph.paragraph_format.alignment)
        if this_title_alignment == 'LEFT (0)':
            this_title_alignment = '左对齐'
        elif this_title_alignment == 'CENTER (1)':
            this_title_alignment = '居中'
        elif this_title_alignment == 'RIGHT (2)':
            this_title_alignment = '右对齐'
        elif this_title_alignment == 'JUSTIFY (3)':
            this_title_alignment = '两端对齐'
        # 与期望的标题信息进行比较并添加错误信息
        expected_this_title_info = RULES[style_name]
        if (this_title_font_name, this_title_font_size, this_title_alignment) != (
                expected_this_title_info['字体'], expected_this_title_info['字号'],expected_this_title_info['对齐方式']):
            error = f"{style_name}的字体为{this_title_font_name}，字号为{this_title_font_size}，" \
                    f"对齐方式为{this_title_alignment}，" \
                    f"不符合规定的{expected_this_title_info['字体']}、{expected_this_title_info['字号']}、" \
                    f"{expected_this_title_info['对齐方式']}"
            errors.append(error)
print("题目和各级标题的段落索引为",title_index)
# 提取正文信息
for j,paragraph in enumerate(doc.paragraphs):
  if j in title_index:
    continue
  else:
    # print(f"第{j}段")
    expected_font_name = RULES['正文']['字体']
    expected_font_size = RULES['正文']['字号']
    expected_indent = RULES['正文']['首行缩进']
    expected_spacing = RULES['正文']['行距']
    this_paragraph = doc.paragraphs[j]
    this_run = this_paragraph.runs[0]
    this_font_name = this_run.font.name
    this_font_size = this_run.font.size.pt
    # print(this_font_name)
    # print(this_font_size)
    this_paragraph_indent = this_paragraph.paragraph_format.first_line_indent.pt
    this_paragraph_spacing = this_paragraph.paragraph_format.line_spacing.pt
    # print(this_paragraph_indent)
    # print(this_paragraph_spacing)
    if (this_font_name, this_font_size, this_paragraph_indent, this_paragraph_spacing) != (
            expected_font_name, expected_font_size, expected_indent, expected_spacing):
        error = f"第{j+1}段的字体为{this_font_name}，字号为{this_font_size}，首行缩进为{this_paragraph_indent}，行距为{this_paragraph_spacing}，不符合规定的{expected_font_name}、{expected_font_size}、{expected_indent}、{expected_spacing}"
        errors.append(error)

# 修正错误信息
if errors:
    print("以下是检测到的格式错误：")
    for error in errors:
        print(error)
#-------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# 执行修正操作
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# 提取页面大小规定
expected_page_size = RULES['页面大小']

# 设置页面大小
sections = doc.sections
if sections:
  for i,section in enumerate(sections):
    section = sections[i]
    page_width = Cm(expected_page_size[1])  # 将宽度转换为Cm单位
    page_height = Cm(expected_page_size[2])  # 将高度转换为Cm单位
    section.page_width = page_width
    section.page_height = page_height

# 设置页边距规格
expected_margins = RULES['页边距']
# 设置页边距
sections = doc.sections
if sections:
  for i,section in enumerate(sections):
    section = sections[i]
    # 左边距
    if round(section.left_margin.cm, 2) != expected_margins['左边距']:
        section.left_margin = Cm(expected_margins['左边距'])
    # 右边距
    if round(section.right_margin.cm, 2) != expected_margins['右边距']:
        section.right_margin = Cm(expected_margins['右边距'])
    # 上边距
    if round(section.top_margin.cm, 2) != expected_margins['上边距']:
        section.top_margin = Cm(expected_margins['上边距'])
    # 下边距
    if round(section.bottom_margin.cm, 2) != expected_margins['下边距']:
        section.bottom_margin = Cm(expected_margins['下边距'])
# 设置题目样式
title_paragraph = doc.paragraphs[0]
expected_title_info = RULES['题目']
# 设置题目字体
title_run = title_paragraph.runs[0]
title_run.font.name = expected_title_info['字体']
title_run.font.size = Pt(expected_title_info['字号'])
# 设置对齐方式
title_alignment = expected_title_info['对齐方式']
if title_alignment == '居中':
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
elif title_alignment == '右对齐':
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
elif title_alignment == '两端对齐':
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
elif title_alignment == '左对齐':
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#设置各级标题
# 遍历文档中的段落
for i, paragraph in enumerate(doc.paragraphs):
    style_name = paragraph.style.name
    # 判断段落的样式是否为各级标题样式
    if style_name in heading_styles:
        # 提取本次标题的字体、字号、对齐方式和行距
        this_title_paragraph = doc.paragraphs[i]
        this_title_run = this_title_paragraph.runs[0]
        this_title_run.font.name = RULES[style_name]['字体']
        this_title_run.font.size = RULES[style_name]['字号']
        # 提取对齐方式
        if RULES[style_name]['对齐方式'] == '左对齐':
            this_title_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif RULES[style_name]['对齐方式'] == '居中':
            this_title_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif RULES[style_name]['对齐方式'] == '右对齐':
            this_title_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif RULES[style_name]['对齐方式'] == '两端对齐':
            this_title_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#设置正文样式
for j,paragraph in enumerate(doc.paragraphs):
  if j in title_index:
    continue
  else:
    # print(f"第{j}段")
    expected_font_name = RULES['正文']['字体']
    expected_font_size = RULES['正文']['字号']
    expected_indent = RULES['正文']['首行缩进']
    expected_spacing = RULES['正文']['行距']
    this_paragraph = doc.paragraphs[j]
    this_run = this_paragraph.runs[0]
    this_run.font.name = expected_font_name
    this_run.font.size = expected_font_size
    this_paragraph.paragraph_format.first_line_indent = expected_indent
    this_paragraph.paragraph_format.line_spacing = expected_spacing
# 保存修正后的文档
#doc.save('your_modified_document.docx')
#------------------------------------------------------------------------------------------
#布局
# 初始化对话列表
dialogue = []

# 定义对话界面布局
st.title("问答界面")

# 在底部输入框获取用户问题输入
user_question = st.text_input("用户问题")

# 当用户发送问题时，添加到对话列表并进行回答
if user_question:
    # 添加用户问题和系统回答到对话列表
    dialogue.append(("用户", user_question))
    # 在这里添加系统回答的逻辑
    system_reply = "系统回答内容"
    dialogue.append(("系统", system_reply))

# 显示对话列表
for role, content in dialogue:
    if role == "用户":
        st.text_area(content, key=role, value=content, height=100)
    elif role == "系统":
        st.text_area(errors, key=role, value=content, height=100)
